import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

MD = 'paper_jfds.md'
DOCX = 'paper_jfds.docx'

with open(MD, 'r', encoding='utf-8') as f:
    lines = f.readlines()

def add_formatted(p, text):
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            r.bold = True
        else:
            r = p.add_run(part)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)

def add_table(doc, rows):
    if len(rows) < 3: return
    h = [c.strip() for c in rows[0].split('|')[1:-1]]
    d = []
    for row in rows[2:]:
        cells = [c.strip() for c in row.split('|')[1:-1]]
        if cells: d.append(cells)
    if not d: return
    nc = min(len(h), max(len(r) for r in d))
    t = doc.add_table(rows=1+len(d), cols=nc)
    t.style = 'Light Grid Accent 1'
    for i, hd in enumerate(h[:nc]):
        c = t.rows[0].cells[i]; c.text = ''
        r = c.paragraphs[0].add_run(hd); r.bold = True
        r.font.size = Pt(9); r.font.name = 'Times New Roman'
    for i, rd in enumerate(d):
        for j, ct in enumerate(rd[:nc]):
            cell = t.rows[i+1].cells[j]; cell.text = ''
            add_formatted(cell.paragraphs[0], ct)

doc = Document()
s = doc.styles['Normal']
s.font.name = 'Times New Roman'; s.font.size = Pt(11)
s.paragraph_format.space_after = Pt(4); s.paragraph_format.line_spacing = 1.15

for sec in doc.sections:
    sec.top_margin = Cm(2.54); sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.54); sec.right_margin = Cm(2.54)

in_tbl = False; tbl_rows = []; title_done = False

for raw in lines:
    row = raw.rstrip(); s = row.strip()
    if not s:
        if in_tbl: add_table(doc, tbl_rows); tbl_rows = []; in_tbl = False
        continue
    if s.startswith('# ') and not title_done:
        title_done = True
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(s[2:]); r.bold = True; r.font.size = Pt(16); r.font.name = 'Times New Roman'
        continue
    if s.startswith('## ') and not s.startswith('### '):
        doc.add_heading(s[3:], level=1); continue
    if s.startswith('### '):
        doc.add_heading(s[4:], level=2); continue
    if s.startswith('|') and s.endswith('|'):
        if not in_tbl: in_tbl = True; tbl_rows = []
        if not re.match(r'^\|[\s\-:|]+\|$', s): tbl_rows.append(row)
        continue
    if in_tbl: add_table(doc, tbl_rows); tbl_rows = []; in_tbl = False
    if s.startswith('- '):
        p = doc.add_paragraph(style='List Bullet')
        add_formatted(p, s[2:]); continue
    p = doc.add_paragraph(); add_formatted(p, s)

doc.save(DOCX)
print(f'Written: {DOCX}')
